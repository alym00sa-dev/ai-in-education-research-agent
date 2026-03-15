"""Export utilities — Word document and JSON session export."""
import io
import json
import re
from datetime import datetime
from typing import Any, Dict, List, Optional


# ── Word export ───────────────────────────────────────────────────────────────

def _add_markdown_body(doc, text: str):
    """Parse a markdown string into Word paragraphs/headings."""
    from docx.shared import Pt  # noqa: F401 – available on import
    for line in text.split("\n"):
        stripped = line.strip()
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif stripped == "---" or stripped == "___":
            doc.add_paragraph("─" * 60)
        elif stripped:
            clean = re.sub(r"\*{1,2}(.+?)\*{1,2}", r"\1", stripped)
            clean = re.sub(r"_{1,2}(.+?)_{1,2}", r"\1", clean)
            doc.add_paragraph(clean)


def export_report_as_docx(
    research_summary: str,
    session_query: str,
    structured_papers: Optional[List[Dict]] = None,
    selected_columns: Optional[List[Dict]] = None,
    results: Optional[Dict[str, Any]] = None,
) -> bytes:
    """Convert the research report to a richly structured .docx file.

    Sections:
      1. Title block
      2. Report body (markdown → Word styles)
      3. Data Extraction Table (user-selected columns)
      4. Quality Assessment appendix (coverage + hypotheses + causality diagram)

    Args:
        research_summary: Final report markdown.
        session_query: Original user query (used in title block).
        structured_papers: Extracted paper metadata rows.
        selected_columns: List of {key, label} dicts matching the on-screen table.
            Defaults to a standard set if not provided.
        results: Full results dict — supplies qa_assessment, swanson_hypotheses,
            causality_diagram if available.

    Returns:
        Raw .docx bytes for st.download_button.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    _DEFAULT_COLUMNS = [
        {"key": "title",               "label": "Title"},
        {"key": "year",                "label": "Year"},
        {"key": "study_design",        "label": "Study Design"},
        {"key": "outcome",             "label": "Outcome"},
        {"key": "measure",             "label": "Study Measure"},
        {"key": "finding_direction",   "label": "Finding Direction"},
        {"key": "effect_size",         "label": "Effect Size"},
        {"key": "confidence_interval", "label": "Confidence Interval"},
        {"key": "std_deviation",       "label": "Std. Deviation"},
        {"key": "study_size",          "label": "Study Size"},
    ]
    _COL_EXTRACTORS = {
        "title":               lambda p: p.get("title", ""),
        "year":                lambda p: str(p.get("year", "") or ""),
        "study_design":        lambda p: p.get("study_design", ""),
        "population":          lambda p: p.get("population", ""),
        "outcome":             lambda p: p.get("outcome", ""),
        "measure":             lambda p: p.get("measure", ""),
        "finding_direction":   lambda p: p.get("finding_direction", ""),
        "effect_size":         lambda p: str(p.get("effect_size", "") or ""),
        "confidence_interval": lambda p: p.get("confidence_interval", ""),
        "std_deviation":       lambda p: p.get("std_deviation", ""),
        "study_size":          lambda p: str(p.get("study_size", "") or ""),
    }

    columns = selected_columns or _DEFAULT_COLUMNS
    results = results or {}

    doc = Document()

    # ── 1. Title block ────────────────────────────────────────────────────
    title = doc.add_heading("EDU Deep Research Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Query: {session_query}")
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y %I:%M %p')}")
    doc.add_paragraph("")

    # ── 2. Report body ────────────────────────────────────────────────────
    _add_markdown_body(doc, research_summary)

    # ── 3. Data Extraction Table ──────────────────────────────────────────
    if structured_papers:
        doc.add_page_break()
        doc.add_heading("Data Extraction Table", level=1)
        n_cols = len(columns)
        table = doc.add_table(rows=1, cols=n_cols)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        for i, col in enumerate(columns):
            hdr[i].text = col["label"]
            hdr[i].paragraphs[0].runs[0].bold = True

        for p in structured_papers:
            row = table.add_row().cells
            for i, col in enumerate(columns):
                extractor = _COL_EXTRACTORS.get(col["key"], lambda _: "")
                row[i].text = extractor(p) or ""

    # ── 4. Quality Assessment appendix ───────────────────────────────────
    qa_assessment = results.get("qa_assessment")
    swanson_hypotheses = results.get("swanson_hypotheses") or []
    causality_diagram = results.get("causality_diagram") or ""

    has_qa_content = any([qa_assessment, swanson_hypotheses, causality_diagram])
    if has_qa_content:
        doc.add_page_break()
        doc.add_heading("Quality Assessment", level=1)

        if qa_assessment:
            doc.add_heading("Coverage Assessment", level=2)
            _add_markdown_body(doc, qa_assessment)

        if swanson_hypotheses:
            doc.add_heading("Novel Hypotheses (Swanson ABC)", level=2)
            doc.add_paragraph(
                "The following novel hypotheses were surfaced by chaining A→B and B→C "
                "mechanism pairs across sub-researcher findings. They represent connections "
                "not explicitly stated in any single source."
            )
            for idx, h in enumerate(swanson_hypotheses):
                confidence = h.get("confidence", "Speculative")
                a, b, c = h.get("A", ""), h.get("B", ""), h.get("C", "")
                mech_ab = h.get("mechanism_AB", "")
                mech_bc = h.get("mechanism_BC", "")
                cites_ab = ", ".join(h.get("citations_AB", [])) or "—"
                cites_bc = ", ".join(h.get("citations_BC", [])) or "—"

                doc.add_heading(f"Hypothesis {idx + 1} [{confidence}]", level=3)
                p_chain = doc.add_paragraph()
                p_chain.add_run(f"{a} → {b} → {c}").bold = True
                doc.add_paragraph(f"A→B: {mech_ab} ({cites_ab})")
                doc.add_paragraph(f"B→C: {mech_bc} ({cites_bc})")

        if causality_diagram and "graph" in causality_diagram:
            doc.add_heading("Causality Diagram (Mermaid)", level=2)
            doc.add_paragraph(
                "The diagram below is in Mermaid syntax. Paste it at mermaid.live to render."
            )
            # Render as a monospace paragraph
            p_diag = doc.add_paragraph()
            run = p_diag.add_run(causality_diagram)
            run.font.name = "Courier New"
            run.font.size = Pt(8)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── JSON session export ────────────────────────────────────────────────────────

def export_session_as_json(
    results: Dict[str, Any],
    event_log: Optional[List[Dict[str, Any]]] = None,
) -> str:
    """Export the full session as a human-readable JSON string.

    Includes: session metadata, final report, QA outputs (coverage assessment,
    extraction table, hypotheses, causality diagram), structured papers,
    sources, sub-researcher notes, and a slim event log for traceability.

    Args:
        results: Full results dict from _finalize (or finalize_streamed_research).
        event_log: Streaming event log from the session (optional).

    Returns:
        Pretty-printed JSON string suitable for st.download_button.
    """
    event_log = event_log or []

    # Slim event log: keep type/node/content, drop token events
    slim_log = [
        {"type": e.get("type"), "node": e.get("node"), "content": e.get("content")}
        for e in event_log
        if e.get("type") not in ("token", "done")
    ]

    session_meta = results.get("session") or {}
    payload = {
        "generated_at": datetime.now().isoformat() + "Z",
        "session_id": session_meta.get("session_id", ""),
        "query": session_meta.get("query", ""),
        "model_provider": session_meta.get("model_provider", ""),
        "search_depth": session_meta.get("search_depth", ""),
        "final_report": results.get("research_summary", ""),
        "qa_assessment": results.get("qa_assessment"),
        "extraction_table": results.get("extraction_table"),
        "swanson_hypotheses": results.get("swanson_hypotheses") or [],
        "causality_diagram": results.get("causality_diagram"),
        "structured_papers": results.get("structured_papers") or [],
        "sources": results.get("sources", []),
        "sub_researcher_notes": results.get("sub_researcher_notes") or [],
        "event_log": slim_log,
    }

    return json.dumps(payload, indent=2, ensure_ascii=False, default=str)


# ── Legacy — kept for backward compatibility ──────────────────────────────────

def export_audit_log_as_json(event_log: List[Dict[str, Any]]) -> str:
    """Serialize the streaming event log to a pretty-printed JSON string.

    Prefer export_session_as_json for new call sites — this retains the
    old signature for any existing uses.
    """
    slim_log = [
        {"type": e.get("type"), "node": e.get("node"), "content": e.get("content")}
        for e in event_log
        if e.get("type") not in ("token", "done")
    ]
    return json.dumps(
        {
            "generated_at": datetime.now().isoformat(),
            "event_count": len(event_log),
            "events": slim_log,
        },
        indent=2,
    )
